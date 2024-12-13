import os
import re
from docx import Document
import fitz  # PyMuPDF


def read_word(file_path):
    """读取 Word 文件内容并返回为纯文本，同时统计字数"""
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    word_count = len(text)  # 统计字数
    print(f"word文件: {file_path}_字数_{word_count}")
    return text, word_count


def read_pdf(file_path):
    """读取 PDF 文件内容并返回为纯文本，同时统计字数"""
    doc = fitz.open(file_path)
    text = ""
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)  # 加载页面
        text += page.get_text("text")  # 提取页面文本
    word_count = len(text)  # 统计字数
    print(f"PDF文件: {file_path}_字数_{word_count}")
    # 清理多余的空格和换行符
    text = clean_text(text)
    return text, word_count


def clean_text(text):
    """清理文本中的多余换行符、空格和其他格式问题"""
    # 替换连续的换行符（\n）为单个空格
    text = re.sub(r'\n+', ' ', text)
    # 替换多余的空格为单个空格
    text = re.sub(r'\s+', ' ', text)
    # 去掉文本开头和结尾的空白字符
    text = text.strip()
    return text


def split_text_by_char_count(text, char_limit):
    """将文本按字符数切割"""
    chunks = [text[i:i + char_limit] for i in range(0, len(text), char_limit)]
    return chunks


def save_text_to_word(chunks, output_folder, base_filename, word_count):
    """将切割后的内容保存到指定文件夹中的 Word 文件"""
    # 为每个文件创建一个以 base_filename 命名的文件夹
    folder_path = os.path.join(output_folder, base_filename)
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    for idx, chunk in enumerate(chunks):
        new_doc = Document()
        # 将切割后的内容写入新的 Word 文件
        new_doc.add_paragraph(chunk)
        output_file = os.path.join(folder_path, f"{base_filename}_part_{idx + 1}.docx")
        new_doc.save(output_file)


def save_pdf_as_word(pdf_path, output_folder):
    """处理 PDF 文件，若字数小于 5000，保存为 Word 文件"""
    text, word_count = read_pdf(pdf_path)  # 提取 PDF 内容并计算字数
    if word_count < 5000:
        chunks = split_text_by_char_count(text, char_limit=1000)  # 根据字符数进行切割
        base_filename = os.path.splitext(os.path.basename(pdf_path))[0]  # 获取文件名（不带扩展名）
        save_text_to_word(chunks, output_folder, base_filename, word_count)  # 保存为 Word 文件
    else:
        print(f"PDF文件字数超出 5000 字，跳过处理: {pdf_path}")


def process_files(input_folder, output_folder, char_limit):
    """处理文件夹中的所有 Word 和 PDF 文件，按字数切割并保存，并显示字数"""
    for filename in os.listdir(input_folder):
        file_path = os.path.join(input_folder, filename)
        if filename.endswith(".docx"):
            text, word_count = read_word(file_path)  # 读取 Word 文件并统计字数
            if word_count < 5000:
                chunks = split_text_by_char_count(text, char_limit)  # 切割文本
                base_filename = os.path.splitext(filename)[0]  # 获取文件名（不带扩展名）
                save_text_to_word(chunks, output_folder, base_filename, word_count)  # 保存为 Word 文件
            else:
                print(f"Word文件字数超出 5000 字，跳过处理: {file_path}")
        elif filename.endswith(".pdf"):
            save_pdf_as_word(file_path, output_folder)  # 将 PDF 保存为 Word 文件


# 使用示例
input_folder = "input_folder"  # 存放 Word 和 PDF 文件的文件夹路径
output_folder = "output_folder"  # 保存切割后的文件的文件夹路径

# 确保输出文件夹存在
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

process_files(input_folder, output_folder, char_limit=1000)
